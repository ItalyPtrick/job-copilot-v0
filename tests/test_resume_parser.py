"""简历解析器测试"""
import pytest
from pathlib import Path


FIXTURES_DIR = Path(__file__).parent / "fixtures"


@pytest.fixture(autouse=True)
def setup_fixtures():
    """创建测试夹具目录和文件"""
    FIXTURES_DIR.mkdir(exist_ok=True)
    yield
    # 清理由测试创建的临时文件（保留目录）


def test_parse_txt():
    """解析 TXT 文件，返回去除首尾空白的文本"""
    from app.modules.resume.parser import parse_resume

    test_file = FIXTURES_DIR / "test_resume.txt"
    test_file.write_text("  张三\n学历：本科\n技能：Python, FastAPI  ", encoding="utf-8")

    text = parse_resume(str(test_file))
    assert "张三" in text
    assert "Python" in text
    assert text == text.strip()  # 首尾空白已去除


def test_parse_docx():
    """解析 DOCX 文件，提取段落文本"""
    from app.modules.resume.parser import parse_resume
    from docx import Document

    test_file = FIXTURES_DIR / "test_resume.docx"
    doc = Document()
    doc.add_paragraph("李四")
    doc.add_paragraph("3年 Java 开发经验")
    doc.save(str(test_file))

    text = parse_resume(str(test_file))
    assert "李四" in text
    assert "Java" in text


def test_parse_pdf():
    """解析 PDF 文件，提取文本内容"""
    from app.modules.resume.parser import parse_resume
    import fitz

    test_file = FIXTURES_DIR / "test_resume.pdf"
    doc = fitz.open()
    page = doc.new_page()
    # PyMuPDF 内置字体不支持中文，使用英文测试
    page.insert_text((72, 72), "Zhang San - Python Engineer")
    doc.save(str(test_file))
    doc.close()

    text = parse_resume(str(test_file))
    assert "Zhang San" in text
    assert "Python" in text


def test_unsupported_format():
    """不支持的文件格式抛出 ValueError"""
    from app.modules.resume.parser import parse_resume

    test_file = FIXTURES_DIR / "test_resume.xyz"
    test_file.write_bytes(b"fake content")

    with pytest.raises(ValueError, match="不支持的文件格式"):
        parse_resume(str(test_file))


def test_parsers_dict_keys():
    """PARSERS 字典覆盖 .pdf/.docx/.txt 三种扩展名"""
    from app.modules.resume.parser import PARSERS

    assert set(PARSERS.keys()) == {".pdf", ".docx", ".txt"}
